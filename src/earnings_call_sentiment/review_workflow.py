"""Thin local review workflow for the demo UI.

Reuses the existing deterministic pipeline and optional summary layer.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import UTC, datetime
import json
import tempfile
import re
import shutil
import subprocess
from pathlib import Path
from types import SimpleNamespace
from typing import Any
from zipfile import ZipFile

import pandas as pd

from earnings_call_sentiment import cli as cli_module
from earnings_call_sentiment.pipeline.run import (
    run_pipeline,
    write_sentiment_artifacts,
    write_transcript_artifacts,
)
from earnings_call_sentiment.post_summary import generate_optional_summary
import earnings_call_sentiment.question_shifts as question_shifts
from earnings_call_sentiment.summary_config import SummaryConfig

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover - optional runtime dependency
    DocxDocument = None

DEFAULT_UI_PROMPT = """You are reviewing one earnings call. Stay grounded in the deterministic artifacts only. Identify the clearest guidance changes, the strongest tone-change moments, the evidence snippets that support them, and any places where the evidence is still ambiguous. Prefer conservative language over confident speculation. Do not make live trading claims or claim predictive edge."""

SUPPORTED_MEDIA_SUFFIXES = {
    ".wav",
    ".mp3",
    ".m4a",
    ".aac",
    ".flac",
    ".ogg",
    ".mp4",
    ".mov",
    ".mkv",
    ".webm",
}
SUPPORTED_DOCUMENT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".doc", ".docx"}


@dataclass(frozen=True)
class ReviewRun:
    run_id: str
    cache_dir: Path
    out_dir: Path
    input_dir: Path


def slugify_token(value: str, fallback: str = "run") -> str:
    token = re.sub(r"[^a-z0-9]+", "-", value.strip().lower()).strip("-")
    return token or fallback


def prepare_review_run(
    *,
    repo_root: Path,
    source_label: str,
    cache_base: Path | None = None,
    out_base: Path | None = None,
) -> ReviewRun:
    timestamp = datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")
    slug = slugify_token(source_label)
    run_id = f"{timestamp}-{slug}"
    cache_root = (cache_base or (repo_root / "cache" / "ui_runs")).resolve()
    out_root = (out_base or (repo_root / "outputs" / "ui_runs")).resolve()
    cache_dir = cache_root / run_id
    out_dir = out_root / run_id
    input_dir = out_dir / "inputs"
    cache_dir.mkdir(parents=True, exist_ok=True)
    input_dir.mkdir(parents=True, exist_ok=True)
    return ReviewRun(run_id=run_id, cache_dir=cache_dir, out_dir=out_dir, input_dir=input_dir)


def save_uploaded_file(source_path: Path, target_dir: Path) -> Path:
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / source_path.name
    shutil.copy2(source_path, target_path)
    return target_path


def extract_text_from_document(path: Path) -> str:
    source = path.expanduser().resolve()
    suffix = source.suffix.lower()

    if suffix in {".txt", ".md", ".csv", ".json"}:
        return source.read_text(encoding="utf-8", errors="ignore").strip()

    if suffix == ".docx":
        text = _extract_text_from_docx(source)
        if text:
            return text
        raise RuntimeError(f"Unable to extract text from DOCX: {source}")

    if suffix == ".doc":
        text = _extract_text_from_doc(source)
        if text:
            return text
        raise RuntimeError(
            "Legacy .doc extraction failed. Install a converter such as antiword, "
            "or on macOS ensure textutil is available."
        )

    raise RuntimeError(
        "Unsupported document type. Use one of: "
        + ", ".join(sorted(SUPPORTED_DOCUMENT_SUFFIXES))
    )


def _extract_text_from_docx(path: Path) -> str:
    if DocxDocument is not None:
        doc = DocxDocument(str(path))
        blocks: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return "\n\n".join(blocks).strip()

    try:
        with ZipFile(path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    except Exception:
        return ""
    text = re.sub(r"</w:p>", "\n\n", xml)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _extract_text_from_doc(path: Path) -> str:
    commands: list[list[str]] = []
    if shutil.which("textutil"):
        commands.append(["textutil", "-convert", "txt", "-stdout", str(path)])
    if shutil.which("antiword"):
        commands.append(["antiword", str(path)])
    if shutil.which("soffice"):
        with tempfile.TemporaryDirectory(prefix="ecs-doc-") as temp_dir:
            commands.append(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    temp_dir,
                    str(path),
                ]
            )

    for command in commands:
        proc = subprocess.run(command, check=False, capture_output=True, text=True)
        if proc.returncode != 0:
            continue
        if command[0] == "soffice":
            converted = Path(command[5]) / f"{path.stem}.txt"
            if converted.exists() and converted.stat().st_size > 0:
                return converted.read_text(encoding="utf-8", errors="ignore").strip()
            continue
        if (proc.stdout or "").strip():
            return proc.stdout.strip()
    return ""


def chunk_text_for_review(text: str, max_chars: int = 900) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", normalized) if part.strip()]
    if not paragraphs:
        return []

    chunks: list[str] = []
    for paragraph in paragraphs:
        if len(paragraph) <= max_chars:
            chunks.append(paragraph)
            continue

        sentences = [item.strip() for item in re.split(r"(?<=[.!?])\s+", paragraph) if item.strip()]
        if not sentences:
            sentences = [paragraph]

        current = ""
        for sentence in sentences:
            candidate = sentence if not current else f"{current} {sentence}"
            if len(candidate) <= max_chars:
                current = candidate
                continue
            if current:
                chunks.append(current.strip())
            if len(sentence) <= max_chars:
                current = sentence
                continue
            for idx in range(0, len(sentence), max_chars):
                piece = sentence[idx : idx + max_chars].strip()
                if piece:
                    chunks.append(piece)
            current = ""
        if current:
            chunks.append(current.strip())
    return chunks


def build_segments_from_text(
    text: str,
    *,
    max_chars: int = 900,
    words_per_second: float = 2.5,
) -> list[dict[str, Any]]:
    chunks = chunk_text_for_review(text, max_chars=max_chars)
    segments: list[dict[str, Any]] = []
    cursor = 0.0
    for chunk in chunks:
        word_count = max(1, len(re.findall(r"\b\w+\b", chunk)))
        duration = max(5.0, min(90.0, word_count / max(words_per_second, 0.5)))
        start = round(cursor, 3)
        end = round(cursor + duration, 3)
        segments.append({"start": start, "end": end, "text": chunk})
        cursor = end
    return segments


def run_document_review(
    *,
    text: str,
    review_run: ReviewRun,
    symbol: str,
    event_dt: str,
    prior_guidance_path: str | None,
    tone_change_threshold: float,
    sentiment_model: str,
    sentiment_revision: str,
    question_shift_enabled: bool,
    pre_window_s: float,
    post_window_s: float,
    min_gap_s: float,
    min_chars: int,
    summary_config: SummaryConfig,
    verbose: bool = False,
) -> dict[str, Any]:
    segments = build_segments_from_text(text)
    if not segments:
        raise RuntimeError("No usable text could be extracted from the document.")

    transcript_json, transcript_txt = write_transcript_artifacts(segments, review_run.out_dir)
    artifacts = write_sentiment_artifacts(
        segments=segments,
        output_path=review_run.out_dir,
        verbose=verbose,
        sentiment_model=sentiment_model,
        sentiment_revision=sentiment_revision,
    )
    chunks_scored_df, chunks_scored_csv, chunks_scored_jsonl = write_chunks_scored_artifacts(
        review_run.out_dir,
        artifacts["sentiment_segments"],
    )
    post_paths = _run_postscore(
        chunks_scored_df=chunks_scored_df,
        out_dir=review_run.out_dir,
        prior_guidance_path=prior_guidance_path,
        tone_change_threshold=tone_change_threshold,
        sentiment_model=sentiment_model,
        sentiment_revision=sentiment_revision,
        video_path=None,
        enable_visual=False,
        visual_note=None,
        youtube_url=None,
        cache_dir=review_run.cache_dir,
        verbose=verbose,
    )
    if question_shift_enabled:
        write_question_shift_artifacts(
            segments=segments,
            chunks_scored_df=chunks_scored_df,
            out_dir=review_run.out_dir,
            pre_window_s=pre_window_s,
            post_window_s=post_window_s,
            min_gap_s=min_gap_s,
            min_chars=min_chars,
        )

    run_meta_path = cli_module._write_run_meta(
        out_dir=review_run.out_dir,
        symbol=symbol,
        event_dt=event_dt,
        source_url=None,
    )
    timing_note_path = review_run.out_dir / "document_timing_note.txt"
    timing_note_path.write_text(
        (
            "Document-mode timing is relative only. Segment start/end values are synthetic "
            "offsets derived from text chunk order and estimated reading duration, not from "
            "real media timestamps."
        ),
        encoding="utf-8",
    )
    summary_path = None
    if summary_config.enabled:
        summary_path = generate_optional_summary(
            out_dir=review_run.out_dir,
            config=summary_config,
            verbose=verbose,
        )

    return {
        "run_id": review_run.run_id,
        "source_kind": "document",
        "out_dir": review_run.out_dir,
        "cache_dir": review_run.cache_dir,
        "transcript_json": transcript_json,
        "transcript_txt": transcript_txt,
        "chunks_scored_csv": chunks_scored_csv,
        "chunks_scored_jsonl": chunks_scored_jsonl,
        "run_meta": run_meta_path,
        "document_timing_note": timing_note_path,
        "summary_path": summary_path,
        **artifacts,
        **post_paths,
    }


def run_media_review(
    *,
    review_run: ReviewRun,
    youtube_url: str | None,
    audio_path: Path | None,
    symbol: str,
    event_dt: str,
    audio_format: str,
    model: str,
    device: str,
    compute_type: str,
    chunk_seconds: float,
    vad: bool,
    prior_guidance_path: str | None,
    tone_change_threshold: float,
    sentiment_model: str,
    sentiment_revision: str,
    question_shift_enabled: bool,
    pre_window_s: float,
    post_window_s: float,
    min_gap_s: float,
    min_chars: int,
    summary_config: SummaryConfig,
    verbose: bool = False,
) -> dict[str, Any]:
    result = run_pipeline(
        youtube_url=youtube_url,
        audio_path=(str(audio_path) if audio_path is not None else None),
        cache_dir=str(review_run.cache_dir),
        out_dir=str(review_run.out_dir),
        verbose=verbose,
        audio_format=audio_format,
        model=model,
        device=device,
        compute_type=compute_type,
        chunk_seconds=chunk_seconds,
        vad=vad,
        sentiment_model=sentiment_model,
        sentiment_revision=sentiment_revision,
    )
    sentiment_segments_path = Path(str(result["sentiment_segments_csv"]))
    sentiment_df = pd.read_csv(sentiment_segments_path)
    chunks_scored_df, chunks_scored_csv, chunks_scored_jsonl = write_chunks_scored_artifacts(
        review_run.out_dir,
        sentiment_df.to_dict(orient="records"),
    )
    post_paths = _run_postscore(
        chunks_scored_df=chunks_scored_df,
        out_dir=review_run.out_dir,
        prior_guidance_path=prior_guidance_path,
        tone_change_threshold=tone_change_threshold,
        sentiment_model=sentiment_model,
        sentiment_revision=sentiment_revision,
        video_path=(
            audio_path.expanduser().resolve()
            if audio_path is not None and cli_module.is_video_path(audio_path)
            else None
        ),
        enable_visual=bool(youtube_url or (audio_path is not None and cli_module.is_video_path(audio_path))),
        visual_note=None if audio_path is not None and cli_module.is_video_path(audio_path) else None,
        youtube_url=youtube_url,
        cache_dir=review_run.cache_dir,
        verbose=verbose,
    )
    if question_shift_enabled:
        segments = json.loads(Path(str(result["transcript_json"])).read_text(encoding="utf-8"))
        write_question_shift_artifacts(
            segments=segments,
            chunks_scored_df=chunks_scored_df,
            out_dir=review_run.out_dir,
            pre_window_s=pre_window_s,
            post_window_s=post_window_s,
            min_gap_s=min_gap_s,
            min_chars=min_chars,
        )

    run_meta_path = cli_module._write_run_meta(
        out_dir=review_run.out_dir,
        symbol=symbol,
        event_dt=event_dt,
        source_url=youtube_url,
    )
    summary_path = None
    if summary_config.enabled:
        summary_path = generate_optional_summary(
            out_dir=review_run.out_dir,
            config=summary_config,
            verbose=verbose,
        )

    return {
        "run_id": review_run.run_id,
        "source_kind": "youtube" if youtube_url else "media",
        "out_dir": review_run.out_dir,
        "cache_dir": review_run.cache_dir,
        "chunks_scored_csv": chunks_scored_csv,
        "chunks_scored_jsonl": chunks_scored_jsonl,
        "run_meta": run_meta_path,
        "summary_path": summary_path,
        **{key: Path(str(value)) if key.endswith(("json", "txt", "csv", "png")) or key == "audio" else value for key, value in result.items()},
        **post_paths,
    }


def write_chunks_scored_artifacts(
    out_dir: Path,
    sentiment_segments: list[dict[str, Any]],
) -> tuple[pd.DataFrame, Path, Path]:
    chunks_scored_df = cli_module._build_chunks_scored_df(sentiment_segments)
    chunks_scored_csv = out_dir / "chunks_scored.csv"
    chunks_scored_df.to_csv(chunks_scored_csv, index=False)
    chunks_scored_jsonl = cli_module._write_chunks_scored_jsonl(chunks_scored_df, out_dir)
    return chunks_scored_df, chunks_scored_csv, chunks_scored_jsonl


def write_question_shift_artifacts(
    *,
    segments: list[dict[str, Any]],
    chunks_scored_df: pd.DataFrame,
    out_dir: Path,
    pre_window_s: float,
    post_window_s: float,
    min_gap_s: float,
    min_chars: int,
) -> Path:
    question_shifts.run(
        segments=segments,
        chunks_scored=chunks_scored_df,
        out_dir=out_dir,
        pre_window_s=pre_window_s,
        post_window_s=post_window_s,
        min_gap_s=min_gap_s,
        min_chars=min_chars,
    )
    return out_dir / "question_sentiment_shifts.csv"


def _run_postscore(
    *,
    chunks_scored_df: pd.DataFrame,
    out_dir: Path,
    prior_guidance_path: str | None,
    tone_change_threshold: float,
    sentiment_model: str,
    sentiment_revision: str,
    video_path: Path | None,
    enable_visual: bool,
    visual_note: str | None,
    youtube_url: str | None,
    cache_dir: Path,
    verbose: bool,
) -> dict[str, Path]:
    resolved_video_path = video_path
    resolved_visual_note = visual_note
    if enable_visual and resolved_video_path is None and youtube_url:
        resolved_video_path, resolved_visual_note = cli_module._resolve_visual_source_path(
            youtube_url=youtube_url,
            audio_path=None,
            cache_dir=cache_dir,
            verbose=verbose,
        )
    args = SimpleNamespace(
        resume=False,
        force=True,
        prior_guidance=prior_guidance_path,
        tone_change_threshold=float(tone_change_threshold),
        sentiment_model=sentiment_model,
        sentiment_revision=sentiment_revision,
    )
    return cli_module._run_postscore_stages(
        chunks_scored_df=chunks_scored_df,
        out_dir=out_dir,
        args=args,
        video_path=resolved_video_path,
        enable_visual=enable_visual,
        visual_note=resolved_visual_note,
    )


def load_artifact_bundle(review_run: ReviewRun) -> dict[str, Any]:
    return load_artifact_bundle_for_dir(
        out_dir=review_run.out_dir,
        run_id=review_run.run_id,
        cache_dir=review_run.cache_dir,
    )


def load_artifact_bundle_for_dir(
    *,
    out_dir: Path,
    run_id: str,
    cache_dir: Path | None = None,
) -> dict[str, Any]:
    out_dir = Path(out_dir).expanduser().resolve()
    bundle: dict[str, Any] = {
        "run_id": run_id,
        "out_dir": str(out_dir),
        "cache_dir": str(cache_dir) if cache_dir is not None else "",
        "artifacts": {},
        "tables": {},
        "json": {},
        "text": {},
    }

    artifact_names = [
        "transcript.txt",
        "transcript.json",
        "document_timing_note.txt",
        "sentiment_segments.csv",
        "sentiment_timeline.png",
        "risk_metrics.json",
        "chunks_scored.csv",
        "chunks_scored.jsonl",
        "guidance.csv",
        "guidance_revision.csv",
        "tone_changes.csv",
        "qa_shift_segments.csv",
        "visual_behavior_frames.csv",
        "visual_behavior_segments.csv",
        "question_sentiment_shifts.csv",
        "question_shifts.png",
        "qa_shift_summary.json",
        "visual_behavior_summary.json",
        "metrics.json",
        "report.md",
        "run_meta.json",
        "llm_summary.json",
    ]
    for name in artifact_names:
        path = out_dir / name
        if path.exists() and path.is_file() and path.stat().st_size > 0:
            bundle["artifacts"][name] = str(path)

    for name in [
        "guidance.csv",
        "guidance_revision.csv",
        "tone_changes.csv",
        "qa_shift_segments.csv",
        "visual_behavior_segments.csv",
        "question_sentiment_shifts.csv",
        "sentiment_segments.csv",
    ]:
        path = out_dir / name
        if path.exists() and path.stat().st_size > 0:
            frame = pd.read_csv(path)
            bundle["tables"][name] = frame.head(12).fillna("").to_dict(orient="records")

    for name in ["metrics.json", "risk_metrics.json", "run_meta.json", "llm_summary.json", "qa_shift_summary.json", "visual_behavior_summary.json"]:
        path = out_dir / name
        if path.exists() and path.stat().st_size > 0:
            bundle["json"][name] = json.loads(path.read_text(encoding="utf-8"))

    for name in ["transcript.txt", "report.md", "document_timing_note.txt"]:
        path = out_dir / name
        if path.exists() and path.stat().st_size > 0:
            bundle["text"][name] = path.read_text(encoding="utf-8")[:12000]

    return bundle
